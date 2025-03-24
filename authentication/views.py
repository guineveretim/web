
from django import forms
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.models import User
from django.shortcuts import render, redirect
from django.urls import reverse
from django.views.decorators.cache import never_cache
from django.shortcuts import render, get_object_or_404

from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from openpyxl import Workbook
from io import BytesIO

from alma import settings
from .models import Application, Status
from django.contrib.auth.models import Group
from .forms import (
    CertificateEducationalDocumentForm,
    DocumentUploadForm,
    EducationalBackgroundCertificateForm,
    EducationalBackgroundMastersForm,
    MastersEducationalDocumentForm,
    ProgramTypeForm,
    StatusForm,
    UserLoginForm,
    UserRegistrationForm,
    PersonalInfoForm,
    ContactAndAddressForm,
    WorkExperienceForm,
    OtherInfoForm,
    
)
from .models import Application, Documents, Other, WorkExperience,Contact
from datetime import datetime
from paynow import Paynow


class PaynowView:
    def __init__(self):
        self.paynow = Paynow(
            settings.PAYNOW['INTEGRATION_ID'],
            settings.PAYNOW['INTEGRATION_KEY'],
            settings.PAYNOW['RESULT_URL'],
            settings.PAYNOW['RETURN_URL']
        )

    def create_payment(self, request):
        payment = self.paynow.create_payment('Application Fee', request.user.email)
        payment.add('Application Fee Payment', 20)
        return payment

    def send_payment(self, payment):
        response = self.paynow.send(payment)
        return response

    def check_transaction_status(self, poll_url):
        status = self.paynow.check_transaction_status(poll_url)
        return status


def index(request):
    return render(request, 'index.html')

def application_list(request):
    applications = Application.objects.all()  # Fetch all applications
    return render(request, 'application/application_list.html', {'applications': applications})



def application_detail(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    return render(request, 'application/application_detail.html', {'application': application})



@login_required
def homepage(request):
    return render(request, 'homepage.html')

@login_required
def logout_view(request):
    logout(request)
    messages.success(request, "You have successfully logged out.")
    return redirect('login')


def register(request):
    if request.method == 'POST':
        form = UserRegistrationForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.set_password(form.cleaned_data['password'])  
            user.save()
             
            group =Group.objects.get(name = 'Students') 
            user.groups.add(group)
            

            login(request, user)  
            messages.success(request, "Registration successful.")
            return redirect('login')  
    else:
        form = UserRegistrationForm()

    return render(request, 'authentication/register.html', {'form': form})

def login_view(request):
    if request.method == 'POST':
        form = UserLoginForm(data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(request, username=username, password=password)

            if user is not None:
            
                if user.groups.filter(name='Registrar').exists():
                    login(request, user)
                    return redirect('application_list')
                elif user.groups.filter(name='Students').exists():
                    login(request, user)
                    return redirect('homepage')
            else:
                
                messages.error(request, 'Invalid username or password.')
        else:
            # Optional: Handle field errors
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, error)

    else:
        form = UserLoginForm()

    return render(request, 'authentication/login.html', {'form': form})


@login_required
def application_view(request):
    user = request.user
    step = request.GET.get('step', '1')
    application, _ = Application.objects.get_or_create(user=user)

    # Step handlers
    step_handlers = {
        '1': step_one_handler,
        '2': step_two_handler,
        '3': step_three_handler,
        '4': step_four_handler,
        '5': step_five_handler,
        '6': step_six_handler,
        '7': step_seven_handler,
        '8': step_eight_handler,
        '9': payment_handler,
        '10': step_nine_handler,
    }

    if step in step_handlers:
        return step_handlers[step](request, application)
    else:  # Redirect if step is invalid
        return redirect(reverse('application_view') + '?step=1')

def step_one_handler(request, application):
    form = ProgramTypeForm(request.POST or None)  # Use POST data if available
    if request.method == 'POST':
        if form.is_valid():
            application.application_type = form.cleaned_data['application_type']
            application.save()
            return redirect(reverse('application_view') + '?step=2')
    return render(request, 'application/program_selection.html', {'form': form})

def step_two_handler(request, application):
    form = PersonalInfoForm(request.POST or None, instance=application)
    if request.method == 'POST' and form.is_valid():
        form.save()
        return redirect(reverse('application_view') + '?step=3')
    return render(request, 'application/personal.html', {'form': form})

def step_three_handler(request, application):
    documents, _ = Documents.objects.get_or_create(application=application)

    form = DocumentUploadForm(request.POST or None, request.FILES or None, instance=documents)
    if request.method == 'POST' and form.is_valid():
        documents = form.save()
        return redirect(reverse('application_view') + '?step=4')

    return render(request, 'application/personal_document.html', {'form': form})

def step_four_handler(request, application):
    contact, _ = Contact.objects.get_or_create(application=application)

    form = ContactAndAddressForm(request.POST or None, instance=contact)
    if request.method == 'POST' and form.is_valid():
        form.save()
        return redirect(reverse('application_view') + '?step=5')

    return render(request, 'application/contact_and_address.html', {'form': form})

def step_five_handler(request, application):
    form_class = EducationalBackgroundMastersForm if application.application_type == 'masters' else EducationalBackgroundCertificateForm
    
    # existing_instance = application.educational.first() if application.educational.exists() else None
    
    template_name = 'application/masters_educational_background.html' if application.application_type == 'masters' else 'application/certificate_educational_background.html'
    existing_instance = application.educational.first() if application.educational.exists() else None
        
    form = form_class(request.POST or None, instance=existing_instance)
    if request.method == 'POST' and form.is_valid():
        educational = form.save(commit=False)
        educational.application = application
        educational.save()
        return redirect(reverse('application_view') + '?step=6')

    print(form.errors)
    return render(request, template_name, {'form': form})

def step_six_handler(request, application):
    if application.application_type == 'masters':
        form_class = MastersEducationalDocumentForm 
        template_name = 'application/mastersdoc.html'
        existing_instance = application.masters.first() if application.masters.exists() else None
    else:  # Certificates
        form_class = CertificateEducationalDocumentForm  
        template_name = 'application/certificatedoc.html'
        existing_instance = application.certificate.first() if application.certificate.exists() else None

    form = form_class(request.POST or None, request.FILES or None, instance=existing_instance)
    if request.method == 'POST' and form.is_valid():
        educational = form.save(commit=False)
        educational.application = application
        educational.save()
        return redirect(reverse('application_view') + '?step=7')

    return render(request, template_name, {'form': form})

def step_seven_handler(request, application):
    WorkExperienceFormSet = forms.inlineformset_factory(Application, WorkExperience, form=WorkExperienceForm, extra=1, can_delete=False)
    formset = WorkExperienceFormSet(request.POST or None, queryset=application.work_experiences.all())
    if request.method == 'POST' and formset.is_valid():
        instances = formset.save(commit=False)
        for instance in instances:
            instance.application = application
            instance.save()
        return redirect(reverse('application_view') + '?step=8')

    return render(request, 'application/work_experience.html', {'formset': formset})

def step_eight_handler(request, application):
    other, _ = Other.objects.get_or_create(application=application)

    form = OtherInfoForm(request.POST or None, instance=other)
    if request.method == 'POST':
        if form.is_valid():
            other_info = form.save()  # No need to commit=False then save again.  form.save() handles everything

            return redirect(reverse('application_view') + '?step=9')
        else:
            print(form.errors)  # Print form errors to the console for debugging

    return render(request, 'application/other.html', {'form': form})

# Payment handling code here remains unchanged
def check_payment_status(paynow, response, application):
    """Check the payment status using the poll URL."""
    status = paynow.check_transaction_status(response.poll_url)
    if status.paid:
        application.is_paid = True
        application.save()
        return True
    return False


def payment_handler(request, application):
    paynow = Paynow(
        settings.PAYNOW['INTEGRATION_ID'],
        settings.PAYNOW['INTEGRATION_KEY'],
        settings.PAYNOW['RESULT_URL'],
        settings.PAYNOW['RETURN_URL']
    )
    context = {  # Define context here
        'application': application  # Optionally pass the application object to the template
    }
    if request.method == 'POST':
        payment_method = request.POST.get('payment_method')
        payment = paynow.create_payment('Application Fee', 'moses.chundu@cccsea.org')
        payment.add('Application Fee Payment', 20)

        if payment_method == 'mobile':
            response = paynow.send_mobile(payment, '0771111111', 'ecocash')
            print(f"Mobile Payment Response: {response.__dict__}")

            if response.success:
                payment_successful = check_payment_status(paynow, response, application) # Call method to change application.is_paid

                if payment_successful:
                    messages.success(request, 'Payment successful.')
                    return redirect('application_success')
            else:
                messages.error(request, 'Mobile payment failed. Please try again.')

        elif payment_method == 'card':
            response = paynow.send(payment)
            print(f"Card Payment Response: {response.__dict__}")

            if response.success:
                if response.has_redirect:
                    messages.info(request, 'Payment is still pending. Please complete the payment.')
                    return redirect(response.redirect_url)
                else:
                    messages.error(request, 'Card payment didn\'t redirect. Please contact support.')
            else:
                messages.error(request, 'Card payment failed. Please try again.')

    return render(request, 'application/payment.html', context)

def step_nine_handler(request, application):
    return render(request, 'application/success.html', {
        'current_year': datetime.now().year,
        'email': request.user.email,
    })
def registrar_view(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    
    status_instance = Status.objects.filter(application=application).order_by('-updated_at').first()

    if status_instance is None:
        
        status_instance = Status(application=application, status='Pending')
        status_instance.save()

    if request.method == "POST":
        form = StatusForm(request.POST, instance=status_instance)
        if form.is_valid():
            form.save() 
            return redirect('application_detail', application_id=application_id)  
    else:
        form = StatusForm(instance=status_instance)

    
    return render(request, 'application/application_detail.html', {
        'form': form,
        'application': application,
        'status': status_instance.status,  
        'last_update': status_instance.last_update  
    })
    
    
@login_required 
def update_status(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    
    
    status, created = Status.objects.get_or_create(application=application)
    
    if request.method == 'POST':
        form = StatusForm(request.POST, instance=status)
        if form.is_valid():
            form.save()  
            return redirect('application_detail', application_id=application.id) 
    else:
        form = StatusForm(instance=status)

    return render(request, 'application/update_status.html', {'form': form, 'application': application})

def send_email(application,status):
    subject = 'Status Update'
    message = f'Your status has been updated to: {status.status}'
    from_email = settings.Email_HOST_USER
    to_email = application.user.email
    send_email(subject, message, from_email, [to_email], fail_silently=False)

def track_application(request):
    application = None
    latest_status = None
    error_message = None
    
    if request.method == 'POST':
        national_id = request.POST.get('national_id')
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        
        try:
            
            application = Application.objects.get(
                national_id=national_id,
                first_name__iexact=first_name,
                last_name__iexact=last_name
            )
            
            latest_status = application.statuses.order_by('-id').first()  
        except Application.DoesNotExist:
            error_message = "No application found with the provided details."

    
    return render(request, 'application/track_application.html', {
        'application': application,
        'latest_status': latest_status,
        'error_message': error_message
    })
    
    
    
def download_pdf(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    contact = application.contacts.first()  # Assuming there's at least one contact related
    masters_edu = application.educational.first()  # Assuming the first educational record for masters
    certificate_edu = application.certificate_edu.first()  # Assuming the first educational record for certificate
    work_experience = application.work_experiences.first()  # Fetch first work experience for simplicity
    documents = application.documents.first()  # Assuming one set of documents is related

    # Create a PDF document
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    p.drawString(100, height - 40, f"Application Detail for {application.first_name} {application.last_name}")
    p.drawString(100, height - 60, f"Application Type: {application.get_application_type_display()}")
    p.drawString(100, height - 80, f"First Name: {application.first_name}")
    p.drawString(100, height - 100, f"Last Name: {application.last_name}")
    p.drawString(100, height - 120, f"Date Of Birth: {application.date_of_birth}")
    p.drawString(100, height - 140, f"Gender: {application.gender}")
    p.drawString(100, height - 160, f"National ID: {application.national_id}")
    p.drawString(100, height - 180, f"Title: {application.get_title_display()}")
    p.drawString(100, height - 200, f"Marital Status: {application.get_marital_status_display()}")
    p.drawString(100, height - 220, f"Place of Birth: {application.place_of_birth}")
    p.drawString(100, height - 240, f"Citizenship: {application.citizenship}")
    p.drawString(100, height - 260, f"Country: {application.country}")
    p.drawString(100, height - 280, f"Disability: {application.disability}")
    p.drawString(100, height - 300, f"Disability Details: {application.disability_details}")
    
    if contact:
        p.drawString(100, height - 320, f"Email: {contact.email}")
        p.drawString(100, height - 340, f"Phone Number: {contact.phone_number}")
        p.drawString(100, height - 360, f"Permanent Address: {contact.permanent_address_no_street}")
        if contact.permanent_address_apt_unit:
            p.drawString(100, height - 380, f"Apt/Unit: {contact.permanent_address_apt_unit}")
        p.drawString(100, height - 400, f"State/Province: {contact.permanent_address_state_province}")

    if masters_edu:
        p.drawString(100, height - 420, f"Programme Applied For (Masters): {masters_edu.get_programme_applied_for_display()}")
        p.drawString(100, height - 440, f"Institution: {masters_edu.awarding_institution}")
        p.drawString(100, height - 460, f"Major Subjects: {masters_edu.major_subjects}")
    
    if certificate_edu:
        p.drawString(100, height - 480, f"Programme Applied For (Certificate): {certificate_edu.get_programme_applied_for_display()}")

    if work_experience:
        p.drawString(100, height - 500, f"Employer: {work_experience.employer}")
        p.drawString(100, height - 520, f"Work Details: {work_experience.work_details}")

    if documents:
        p.drawString(100, height - 540, f"Birth Certificate: {documents.birth_certificate.url if documents.birth_certificate else 'Not Provided'}")
        p.drawString(100, height - 560, f"ID Document: {documents.id_document.url if documents.id_document else 'Not Provided'}")

    p.showPage()
    p.save()
    buffer.seek(0)

    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{application.first_name}_{application.last_name}.pdf"'
    return response


def download_word(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    contact = application.contacts.first()  # Grab the first contact if available
    masters_edu = application.educational.first()  # First Masters educational record
    certificate_edu = application.certificate_edu.first()  # First Certificate educational record
    work_experience = application.work_experiences.first()  # First work experience

    # Create a Word document
    doc = Document()
    doc.add_heading(f'Application Detail for {application.first_name} {application.last_name}', level=1)

    doc.add_paragraph(f'Application Type: {application.get_application_type_display()}')
    doc.add_paragraph(f'First Name: {application.first_name}')
    doc.add_paragraph(f'Last Name: {application.last_name}')
    doc.add_paragraph(f'Date Of Birth: {application.date_of_birth}')
    doc.add_paragraph(f'Gender: {application.gender}')
    doc.add_paragraph(f'National ID: {application.national_id}')
    doc.add_paragraph(f'Title: {application.get_title_display()}')
    doc.add_paragraph(f'Marital Status: {application.get_marital_status_display()}')
    doc.add_paragraph(f'Place of Birth: {application.place_of_birth}')
    doc.add_paragraph(f'Citizenship: {application.citizenship}')
    doc.add_paragraph(f'Country: {application.country}')
    doc.add_paragraph(f'Disability: {application.disability}')
    doc.add_paragraph(f'Disability Details: {application.disability_details}')
    
    if contact:
        doc.add_paragraph(f'Email: {contact.email}')
        doc.add_paragraph(f'Phone Number: {contact.phone_number}')
        doc.add_paragraph(f'Permanent Address: {contact.permanent_address_no_street}')
        if contact.permanent_address_apt_unit:
            doc.add_paragraph(f'Apt/Unit: {contact.permanent_address_apt_unit}')
        doc.add_paragraph(f'State/Province: {contact.permanent_address_state_province}')

    if masters_edu:
        doc.add_paragraph(f'Programme Applied For (Masters): {masters_edu.get_programme_applied_for_display()}')
        doc.add_paragraph(f'Institution: {masters_edu.awarding_institution}')
        doc.add_paragraph(f'Major Subjects: {masters_edu.major_subjects}')

    if certificate_edu:
        doc.add_paragraph(f'Programme Applied For (Certificate): {certificate_edu.get_programme_applied_for_display()}')

    if work_experience:
        doc.add_paragraph(f'Employer: {work_experience.employer}')
        doc.add_paragraph(f'Work Details: {work_experience.work_details}')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{application.first_name}_{application.last_name}.docx"'
    return response


def download_excel(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    contact = application.contacts.first()  # First contact record
    masters_edu = application.educational.first()  # First Masters educational record
    certificate_edu = application.certificate_edu.first()  # First Certificate educational record
    work_experience = application.work_experiences.first()  # First work experience

    # Create a Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Application Details"

    headers = [
        "First Name", 
        "Last Name", 
        "Date Of Birth", 
        "Gender", 
        "National ID", 
        "Title", 
        "Marital Status",
        "Place of Birth",
        "Citizenship",
        "Country",
        "Disability",
        "Disability Details",
        "Email",
        "Phone Number",
        "Permanent Address",
        "Apt/Unit",
        "State/Province",
        "Masters Programme",
        "Masters Institution",
        "Masters Major Subjects",
        "Certificate Programme",
        "Work Employer",
        "Work Details",
        "Documents (Birth Certificate)",
        "Documents (ID)"
    ]

    ws.append(headers)

    row = [
        application.first_name,
        application.last_name,
        application.date_of_birth,
        application.gender,
        application.national_id,
        application.get_title_display(),
        application.get_marital_status_display(),
        application.place_of_birth,
        application.citizenship,
        application.country,
        application.disability,
        application.disability_details,
        contact.email if contact else '',
        contact.phone_number if contact else '',
        contact.permanent_address_no_street if contact else '',
        contact.permanent_address_apt_unit if contact else '',
        contact.permanent_address_state_province if contact else '',
        masters_edu.get_programme_applied_for_display() if masters_edu else '',
        masters_edu.awarding_institution if masters_edu else '',
        masters_edu.major_subjects if masters_edu else '',
        certificate_edu.get_programme_applied_for_display() if certificate_edu else '',
        work_experience.employer if work_experience else '',
        work_experience.work_details if work_experience else '',
        application.documents.first().birth_certificate.url if application.documents.exists() and application.documents.first().birth_certificate else 'Not Provided',
        application.documents.first().id_document.url if application.documents.exists() and application.documents.first().id_document else 'Not Provided',
    ]

    ws.append(row)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{application.first_name}_{application.last_name}.xlsx"'

    wb.save(response)
    return response
    
